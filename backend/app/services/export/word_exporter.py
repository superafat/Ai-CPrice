"""
Word 匯出服務 - 將 Markdown+LaTeX 轉換為包含 OMML 公式的 Word 文檔
"""
import os
import asyncio
import subprocess
from typing import Dict, List, Optional
from loguru import logger
import tempfile
import shutil

from ...models.ocr_result import ProcessingTask, ExportSettings, ImageBlock, BlockType
from ...config import WORD_STYLES


class WordExporter:
    """Word 文檔匯出器"""
    
    def __init__(self):
        self.temp_dir = "temp_export"
        os.makedirs(self.temp_dir, exist_ok=True)
        self.pandoc_available = False
    
    async def initialize(self):
        """初始化匯出器"""
        try:
            # 檢查 Pandoc 是否可用
            result = subprocess.run(["pandoc", "--version"], capture_output=True, text=True)
            if result.returncode == 0:
                self.pandoc_available = True
                logger.info("Pandoc 可用，支援 LaTeX → OMML 轉換")
            else:
                logger.warning("Pandoc 不可用，將使用替代方案")
                
        except FileNotFoundError:
            logger.warning("Pandoc 未安裝，將使用替代方案")
            await self._install_pandoc()
    
    async def _install_pandoc(self):
        """安裝 Pandoc"""
        try:
            logger.info("正在安裝 Pandoc...")
            result = subprocess.run([
                "apt-get", "update", "&&", 
                "apt-get", "install", "-y", "pandoc"
            ], shell=True, capture_output=True, text=True)
            
            if result.returncode == 0:
                self.pandoc_available = True
                logger.info("Pandoc 安裝成功")
            else:
                logger.error(f"Pandoc 安裝失敗: {result.stderr}")
                
        except Exception as e:
            logger.error(f"安裝 Pandoc 時發生錯誤: {e}")
    
    async def export_to_word(
        self, 
        task: ProcessingTask, 
        export_settings: ExportSettings
    ) -> str:
        """
        將處理結果匯出為 Word 文檔
        返回生成的 .docx 檔案路徑
        """
        try:
            logger.info(f"開始匯出任務 {task.task_id} 為 Word 文檔")
            
            # 1. 生成標準化 Markdown
            markdown_content = await self._generate_markdown(task, export_settings)
            
            # 2. 建立臨時 Markdown 檔案
            markdown_path = os.path.join(self.temp_dir, f"{task.task_id}.md")
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            # 3. 建立自訂樣式檔案
            reference_docx = await self._create_reference_document(export_settings)
            
            # 4. 使用 Pandoc 轉換為 Word
            output_path = os.path.join("outputs", f"{task.task_id}.docx")
            await self._convert_with_pandoc(markdown_path, output_path, reference_docx)
            
            # 5. 後處理：確保圖片正確錨定
            await self._post_process_word_document(output_path, export_settings)
            
            logger.info(f"Word 文檔匯出完成: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Word 匯出失敗: {e}")
            raise
    
    async def _generate_markdown(self, task: ProcessingTask, settings: ExportSettings) -> str:
        """生成標準化 Markdown 內容"""
        markdown_lines = []
        
        # 按順序處理所有區塊
        sorted_blocks = sorted(task.blocks, key=lambda x: x.order_in_document)
        
        current_problem = None
        current_options = []
        current_figures = []
        
        for block in sorted_blocks:
            if not block.ocr_result:
                continue
            
            text = block.manual_correction or block.ocr_result.text
            
            if block.type == BlockType.PROBLEM:
                # 如果有前一題，先輸出
                if current_problem is not None:
                    markdown_lines.extend(self._format_problem_section(
                        current_problem, current_options, current_figures, settings
                    ))
                
                current_problem = text
                current_options = []
                current_figures = []
                
            elif block.type == BlockType.OPTIONS:
                current_options.append(text)
                
            elif block.type == BlockType.FIGURE:
                current_figures.append({
                    "path": block.image_path,
                    "order": block.order_in_document,
                    "caption": self._generate_figure_caption(block, settings)
                })
                
            elif block.type == BlockType.SOLUTION:
                # 詳解區塊
                markdown_lines.append(f"\n**解析：**\n")
                markdown_lines.append(f"{text}\n")
        
        # 輸出最後一題
        if current_problem is not None:
            markdown_lines.extend(self._format_problem_section(
                current_problem, current_options, current_figures, settings
            ))
        
        return "\n".join(markdown_lines)
    
    def _format_problem_section(
        self, 
        problem_text: str, 
        options: List[str], 
        figures: List[Dict],
        settings: ExportSettings
    ) -> List[str]:
        """格式化題目區段"""
        lines = []
        
        # 題幹
        lines.append(f"\n## 題目\n")
        
        # 插入圖片（如果有）
        inline_figures = [f for f in figures if self._is_inline_figure(f)]
        block_figures = [f for f in figures if not self._is_inline_figure(f)]
        
        # 題幹文字
        lines.append(f"{problem_text}")
        
        # 行內圖片
        for fig in inline_figures:
            lines.append(f"\n![圖片]({fig['path']})")
            if settings.include_figure_captions and fig['caption']:
                lines.append(f"\n*{fig['caption']}*")
        
        # 區塊圖片
        for fig in block_figures:
            lines.append(f"\n![圖片]({fig['path']})")
            if settings.include_figure_captions and fig['caption']:
                lines.append(f"\n*{fig['caption']}*")
        
        # 選項
        if options:
            lines.append(f"\n### 選項\n")
            option_labels = ['(A)', '(B)', '(C)', '(D)', '(E)', '(F)']
            for i, option in enumerate(options):
                if i < len(option_labels):
                    lines.append(f"{option_labels[i]} {option}")
        
        return lines
    
    def _is_inline_figure(self, figure: Dict) -> bool:
        """判斷圖片是否為行內小圖"""
        # 簡化判斷：根據檔案大小或其他特徵
        # 實際應該根據圖片尺寸和上下文位置判斷
        return "inline" in figure.get("caption", "").lower()
    
    def _generate_figure_caption(self, block: ImageBlock, settings: ExportSettings) -> str:
        """生成圖片說明"""
        if not settings.include_figure_captions:
            return ""
        
        # 簡化的圖片說明生成
        return f"Figure {block.order_in_document}"
    
    async def _create_reference_document(self, settings: ExportSettings) -> str:
        """建立參考文檔（定義樣式）"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.style import WD_STYLE_TYPE
            
            # 建立新文檔
            doc = Document()
            
            # 定義段落樣式
            styles = doc.styles
            
            # 題幹樣式
            problem_style = styles.add_style('CCH-Problem', WD_STYLE_TYPE.PARAGRAPH)
            problem_style.font.name = WORD_STYLES["problem"]["font_family"]
            problem_style.font.size = Pt(settings.problem_font_size)
            
            # 選項樣式
            options_style = styles.add_style('CCH-Options', WD_STYLE_TYPE.PARAGRAPH)
            options_style.font.name = WORD_STYLES["options"]["font_family"]
            options_style.font.size = Pt(settings.problem_font_size)
            
            # 詳解樣式
            solution_style = styles.add_style('CCH-Solution', WD_STYLE_TYPE.PARAGRAPH)
            solution_style.font.name = WORD_STYLES["solution"]["font_family"]
            solution_style.font.size = Pt(settings.solution_font_size)
            
            # 儲存參考文檔
            reference_path = os.path.join(self.temp_dir, f"reference_{settings.problem_font_size}_{settings.solution_font_size}.docx")
            doc.save(reference_path)
            
            return reference_path
            
        except ImportError:
            logger.warning("python-docx 不可用，將使用預設樣式")
            return None
        except Exception as e:
            logger.error(f"建立參考文檔失敗: {e}")
            return None
    
    async def _convert_with_pandoc(
        self, 
        markdown_path: str, 
        output_path: str, 
        reference_docx: Optional[str] = None
    ):
        """使用 Pandoc 轉換 Markdown 為 Word"""
        if not self.pandoc_available:
            raise RuntimeError("Pandoc 不可用，無法進行轉換")
        
        try:
            # 建立 Pandoc 命令
            cmd = [
                "pandoc",
                markdown_path,
                "-o", output_path,
                "--from", "markdown",
                "--to", "docx",
                "--mathml",  # 關鍵：將 LaTeX 轉換為 MathML/OMML
                "--standalone"
            ]
            
            # 如果有參考文檔，使用自訂樣式
            if reference_docx and os.path.exists(reference_docx):
                cmd.extend(["--reference-doc", reference_docx])
            
            # 在執行緒中執行轉換
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                None,
                lambda: subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"Pandoc 轉換失敗: {result.stderr}")
            
            logger.info("Pandoc 轉換完成")
            
        except subprocess.TimeoutExpired:
            logger.error("Pandoc 轉換超時")
            raise
        except Exception as e:
            logger.error(f"Pandoc 轉換失敗: {e}")
            raise
    
    async def _post_process_word_document(self, docx_path: str, settings: ExportSettings):
        """後處理 Word 文檔"""
        try:
            from docx import Document
            from docx.shared import Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 開啟文檔
            doc = Document(docx_path)
            
            # 處理圖片錨定
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # 檢查是否包含圖片
                    if run._element.xpath('.//pic:pic'):
                        # 設定圖片為"與文字同列"
                        # 這需要更複雜的 XML 操作
                        pass
            
            # 處理圖片尺寸
            for shape in doc.inline_shapes:
                if shape.width.cm > settings.image_max_width_cm:
                    # 等比例縮放
                    ratio = settings.image_max_width_cm / shape.width.cm
                    shape.width = Cm(settings.image_max_width_cm)
                    shape.height = Cm(shape.height.cm * ratio)
            
            # 儲存修改
            doc.save(docx_path)
            logger.info("Word 文檔後處理完成")
            
        except ImportError:
            logger.warning("python-docx 不可用，跳過後處理")
        except Exception as e:
            logger.error(f"Word 文檔後處理失敗: {e}")
    
    async def validate_omml_formulas(self, docx_path: str) -> Dict[str, Any]:
        """驗證 Word 文檔中的 OMML 公式"""
        try:
            from docx import Document
            import xml.etree.ElementTree as ET
            
            doc = Document(docx_path)
            
            formula_count = 0
            omml_count = 0
            validation_results = []
            
            # 檢查文檔中的數學元素
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # 查找 OMML 數學元素
                    math_elements = run._element.xpath('.//m:oMath', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})
                    
                    for math_element in math_elements:
                        formula_count += 1
                        omml_count += 1
                        
                        # 驗證 OMML 結構
                        is_valid = self._validate_omml_structure(math_element)
                        validation_results.append({
                            "formula_index": formula_count,
                            "is_valid_omml": is_valid,
                            "xml_content": ET.tostring(math_element, encoding='unicode')[:100] + "..."
                        })
            
            return {
                "total_formulas": formula_count,
                "omml_formulas": omml_count,
                "omml_ratio": omml_count / formula_count if formula_count > 0 else 0,
                "all_formulas_omml": omml_count == formula_count,
                "validation_details": validation_results
            }
            
        except ImportError:
            logger.warning("無法驗證 OMML，python-docx 不可用")
            return {"error": "validation_unavailable"}
        except Exception as e:
            logger.error(f"OMML 驗證失敗: {e}")
            return {"error": str(e)}
    
    def _validate_omml_structure(self, math_element) -> bool:
        """驗證 OMML 數學元素結構"""
        try:
            # 檢查是否包含必要的 OMML 元素
            required_namespaces = ['http://schemas.openxmlformats.org/officeDocument/2006/math']
            
            # 簡化驗證：檢查是否有數學內容
            math_content = math_element.text or ""
            return len(math_content.strip()) > 0
            
        except Exception:
            return False
    
    async def batch_export(
        self, 
        tasks: List[ProcessingTask], 
        export_settings: ExportSettings
    ) -> str:
        """批次匯出多個任務為單一 Word 文檔"""
        try:
            logger.info(f"開始批次匯出 {len(tasks)} 個任務")
            
            # 合併所有任務的內容
            combined_markdown = []
            
            for i, task in enumerate(tasks):
                combined_markdown.append(f"# 題目 {i + 1}\n")
                task_markdown = await self._generate_markdown(task, export_settings)
                combined_markdown.append(task_markdown)
                combined_markdown.append("\n---\n")  # 分隔線
            
            # 建立合併的 Markdown 檔案
            batch_id = f"batch_{len(tasks)}_{int(time.time())}"
            markdown_path = os.path.join(self.temp_dir, f"{batch_id}.md")
            
            with open(markdown_path, 'w', encoding='utf-8') as f:
                f.write("\n".join(combined_markdown))
            
            # 轉換為 Word
            output_path = os.path.join("outputs", f"{batch_id}.docx")
            reference_docx = await self._create_reference_document(export_settings)
            await self._convert_with_pandoc(markdown_path, output_path, reference_docx)
            
            logger.info(f"批次匯出完成: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"批次匯出失敗: {e}")
            raise
    
    async def cleanup_temp_files(self, keep_outputs: bool = True):
        """清理臨時檔案"""
        try:
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                os.makedirs(self.temp_dir, exist_ok=True)
                logger.info("臨時檔案已清理")
                
        except Exception as e:
            logger.error(f"清理臨時檔案失敗: {e}")


import time